#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
增强docx文档，使其更加丰富美观
- 为标题添加emoji图标
- 增强格式（颜色、字体、间距等）
- 处理特殊文本（加粗、链接、代码等）
- 添加更多视觉元素
"""

import re
import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn

# 标题图标映射
TITLE_ICONS = {
    "引言": "💡",
    "项目概述": "📋",
    "设计哲学": "🎯",
    "平台定位": "📍",
    "核心价值": "✨",
    "技术架构": "🏗️",
    "模块化设计理念": "🧩",
    "数据流转架构": "🔄",
    "存储方案": "💾",
    "核心AI能力": "🤖",
    "全面的AI技术栈": "🔧",
    "革命性的零样本标注技术": "🚀",
    "多场景预训练模型": "📦",
    "IoT能力": "🌐",
    "完整的设备生命周期管理": "📱",
    "强大的规则引擎": "⚙️",
    "数据智能分析": "📊",
    "部署灵活性": "☁️",
    "独立部署优势": "🔀",
    "一键部署方案": "⚡",
    "应用场景适配": "🎨",
    "核心优势": "⭐",
    "多语言混编架构": "🔤",
    "零样本标注技术": "🎯",
    "灵活部署": "🌍",
    "丰富生态支持": "🌳",
    "持续迭代优化": "🔄",
    "应用场景": "🎬",
    "人群密度管控": "👥",
    "周界防护": "🛡️",
    "跌倒检测": "⚠️",
    "异常逗留识别": "👀",
    "肢体冲突预警": "⚔️",
    "非法闯入检测": "🚫",
    "公共场所控烟": "🚭",
    "人流统计管控": "📈",
    "区域越界预警": "🚧",
    "环境安全检查": "🔍",
    "火灾早预警": "🔥",
    "扩展应用领域": "🔮",
    "系统展示": "🖼️",
    "核心功能界面展示": "💻",
    "技术实现": "💻",
    "设备控制核心逻辑": "🎮",
    "安全认证体系": "🔐",
    "AI模型管理": "🧠",
    "高性能任务处理": "⚡",
    "功能介绍": "📖",
    "设备管理模块": "📱",
    "流媒体管理模块": "🎥",
    "数据标注模块": "✏️",
    "模型训练与管理模块": "🎓",
    "AI智能分析模块": "🔬",
    "规则引擎模块": "⚙️",
    "系统管理模块": "🛠️",
    "数据统计与分析模块": "📊",
    "部署安装": "🚀",
    "部署要求": "📋",
    "部署优势": "✅",
    "社区与开源": "❤️",
    "我们的承诺": "🤝",
    "加入我们": "🌟",
    "演示环境与支持": "🌐",
    "在线演示": "💻",
    "结语": "🎉",
    "联系方式": "📞",
}

def get_icon_for_title(title):
    """为标题获取合适的图标"""
    # 移除可能的emoji和Markdown格式
    clean_title = re.sub(r'[\U0001F300-\U0001F9FF\U0001FA00-\U0001FAFF\U00002600-\U000027BF\U0001F600-\U0001F64F\U0001F680-\U0001F6FF\U0001F1E0-\U0001F1FF]+', '', title).strip()
    clean_title = re.sub(r'\*\*([^*]+)\*\*', r'\1', clean_title).strip()
    
    # 先尝试完整匹配
    if clean_title in TITLE_ICONS:
        return TITLE_ICONS[clean_title]
    
    # 尝试部分匹配
    for key, icon in TITLE_ICONS.items():
        if key in clean_title or clean_title in key:
            return icon
    
    # 根据标题内容智能匹配
    if '模块' in clean_title:
        return '📦'
    elif '管理' in clean_title:
        return '🛠️'
    elif '技术' in clean_title or '架构' in clean_title:
        return '🏗️'
    elif '部署' in clean_title or '安装' in clean_title:
        return '🚀'
    elif '场景' in clean_title or '应用' in clean_title:
        return '🎬'
    elif '能力' in clean_title or '功能' in clean_title:
        return '✨'
    elif '分析' in clean_title or '统计' in clean_title:
        return '📊'
    elif '训练' in clean_title or '模型' in clean_title:
        return '🧠'
    elif '设备' in clean_title:
        return '📱'
    elif '视频' in clean_title or '流媒体' in clean_title:
        return '🎥'
    elif '数据' in clean_title:
        return '💾'
    elif '安全' in clean_title:
        return '🔐'
    elif '规则' in clean_title or '引擎' in clean_title:
        return '⚙️'
    elif '演示' in clean_title or '环境' in clean_title:
        return '🌐'
    elif '社区' in clean_title or '开源' in clean_title:
        return '❤️'
    elif '联系' in clean_title:
        return '📞'
    elif '结语' in clean_title:
        return '🎉'
    elif '概述' in clean_title or '介绍' in clean_title:
        return '📋'
    elif '优势' in clean_title:
        return '⭐'
    elif '价值' in clean_title:
        return '✨'
    else:
        return '📌'

def is_title_paragraph(para):
    """判断段落是否为标题"""
    text = para.text.strip()
    if not text:
        return False
    
    # 检查是否已经是标题样式
    if para.style.name.startswith('Heading'):
        return True
    
    # 检查字体大小（标题通常较大）
    if para.runs:
        for run in para.runs:
            if run.font.size and run.font.size.pt >= 16:
                return True
    
    # 检查是否加粗
    if para.runs:
        all_bold = all(run.bold for run in para.runs if run.text.strip())
        if all_bold and len(text) < 50:
            return True
    
    # 检查是否包含常见标题关键词
    title_keywords = ['概述', '引言', '结语', '模块', '能力', '架构', '方案', '场景', 
                     '优势', '技术', '管理', '系统', '部署', '介绍', '展示', '实现']
    if any(keyword in text for keyword in title_keywords) and len(text) < 50:
        return True
    
    return False

def extract_text_without_emoji(text):
    """提取文本，移除已有emoji"""
    # 移除emoji
    text = re.sub(r'[\U0001F300-\U0001F9FF\U0001FA00-\U0001FAFF\U00002600-\U000027BF\U0001F600-\U0001F64F\U0001F680-\U0001F6FF\U0001F1E0-\U0001F1FF]+', '', text).strip()
    # 移除Markdown加粗标记
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text).strip()
    return text

def process_inline_formatting(para, text):
    """处理行内格式（加粗、代码、链接）"""
    para.clear()
    
    # 处理加粗 **text**
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # 加粗文本
            run = para.add_run(part[2:-2])
            run.font.bold = True
            run.font.color.rgb = RGBColor(231, 76, 60)  # 红色加粗
        elif part.startswith('`') and part.endswith('`'):
            # 代码文本
            run = para.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(52, 152, 219)  # 蓝色
        else:
            # 处理链接 [text](url)
            link_parts = re.split(r'(\[[^\]]+\]\([^\)]+\))', part)
            for link_part in link_parts:
                match = re.match(r'\[([^\]]+)\]\(([^\)]+)\)', link_part)
                if match:
                    link_text = match.group(1)
                    link_url = match.group(2)
                    run = para.add_run(link_text)
                    run.font.color.rgb = RGBColor(52, 152, 219)  # 蓝色
                    run.underline = True
                elif link_part:
                    para.add_run(link_part)

def enrich_docx(input_path, output_path=None):
    """
    增强docx文档，使其更加丰富美观
    
    Args:
        input_path: 输入docx文件路径
        output_path: 输出docx文件路径，如果为None则自动生成
    """
    if not os.path.exists(input_path):
        print(f"错误：文件不存在: {input_path}")
        return False
    
    # 如果没有指定输出路径，自动生成
    if output_path is None:
        base_name = os.path.splitext(input_path)[0]
        output_path = f"{base_name}_增强版.docx"
    
    # 读取文档
    doc = Document(input_path)
    
    # 公众号常用字体
    font_name = "微软雅黑"
    
    # 设置默认样式
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(14)
    
    # 创建新文档
    new_doc = Document()
    
    # 设置新文档的默认样式
    new_style = new_doc.styles['Normal']
    new_font = new_style.font
    new_font.name = font_name
    new_font.size = Pt(14)
    
    # 处理每个段落
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if not text:
            # 空段落，添加空行
            new_doc.add_paragraph()
            continue
        
        # 判断是否为标题
        if is_title_paragraph(para):
            # 提取标题文本（移除已有emoji）
            clean_title = extract_text_without_emoji(text)
            
            # 获取图标
            icon = get_icon_for_title(clean_title)
            
            # 创建标题段落
            title_para = new_doc.add_paragraph()
            title_para_format = title_para.paragraph_format
            title_para_format.space_before = Pt(16)  # 标题前间距
            title_para_format.space_after = Pt(12)  # 标题后间距
            title_para_format.line_spacing = 1.3
            
            # 添加图标
            icon_run = title_para.add_run(icon + " ")
            icon_run.font.name = font_name
            icon_run.font.size = Pt(20)
            
            # 添加标题文本
            title_run = title_para.add_run(clean_title)
            title_run.font.name = font_name
            title_run.font.size = Pt(20)  # 标题20pt
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(0, 102, 204)  # 蓝色
            
        elif text.startswith('项目地址') and 'http' in text:
            # 项目地址特殊处理
            url_para = new_doc.add_paragraph()
            url_para_format = url_para.paragraph_format
            url_para_format.space_after = Pt(8)
            
            # 分离"项目地址："和URL
            match = re.search(r'(项目地址[：:]\s*)(https?://[^\s]+)', text)
            if match:
                prefix = match.group(1)
                url = match.group(2)
                
                # 添加前缀
                prefix_run = url_para.add_run(prefix)
                prefix_run.font.name = font_name
                prefix_run.font.size = Pt(12)
                
                # 添加URL（红色、小字体）
                url_run = url_para.add_run(url)
                url_run.font.name = font_name
                url_run.font.size = Pt(11)
                url_run.font.color.rgb = RGBColor(220, 20, 60)  # 红色
                url_run.underline = True
            else:
                # 如果匹配失败，整个文本设为红色小字体
                run = url_para.add_run(text)
                run.font.name = font_name
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(220, 20, 60)
                
        elif text.strip() == '---' or text.strip() == '——':
            # 分隔线
            hr_para = new_doc.add_paragraph('─' * 50)
            hr_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            hr_para_format = hr_para.paragraph_format
            hr_para_format.space_before = Pt(12)
            hr_para_format.space_after = Pt(12)
            for run in hr_para.runs:
                run.font.color.rgb = RGBColor(200, 200, 200)  # 灰色
                run.font.size = Pt(10)
                
        elif text.startswith('> '):
            # 引用块
            quote_text = text[2:].strip()
            quote_para = new_doc.add_paragraph()
            quote_para_format = quote_para.paragraph_format
            quote_para_format.left_indent = Inches(0.3)
            quote_para_format.space_before = Pt(8)
            quote_para_format.space_after = Pt(8)
            
            # 添加引用标记
            quote_marker = quote_para.add_run("💬 ")
            quote_marker.font.size = Pt(14)
            
            # 添加引用文本
            quote_run = quote_para.add_run(quote_text)
            quote_run.font.italic = True
            quote_run.font.color.rgb = RGBColor(102, 102, 102)  # 灰色
            quote_run.font.size = Pt(13)
            
        else:
            # 普通段落
            content_para = new_doc.add_paragraph()
            content_para_format = content_para.paragraph_format
            content_para_format.space_after = Pt(8)
            content_para_format.line_spacing = 1.6  # 行距1.6倍
            
            # 检查是否包含特殊格式
            if '**' in text or '`' in text or '[' in text and '](' in text:
                # 包含Markdown格式，需要特殊处理
                process_inline_formatting(content_para, text)
                # 设置默认字体
                for run in content_para.runs:
                    if not run.font.name or run.font.name == 'Calibri':
                        run.font.name = font_name
                    if not run.font.size:
                        run.font.size = Pt(14)
            else:
                # 普通文本，检查是否包含项目地址
                if '项目地址' in text and 'http' in text:
                    # 分段处理：项目地址部分用红色小字体，其他部分正常
                    url_pattern = r'(项目地址[：:]\s*)(https?://[^\s]+)'
                    parts = re.split(url_pattern, text)
                    for part in parts:
                        if not part:
                            continue
                        if re.match(r'https?://', part):
                            # URL部分：红色、小字体
                            run = content_para.add_run(part)
                            run.font.name = font_name
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(220, 20, 60)
                            run.underline = True
                        elif part.startswith('项目地址'):
                            # "项目地址："前缀
                            run = content_para.add_run(part)
                            run.font.name = font_name
                            run.font.size = Pt(12)
                        else:
                            # 其他文本：正常格式
                            run = content_para.add_run(part)
                            run.font.name = font_name
                            run.font.size = Pt(14)
                else:
                    # 普通正文
                    run = content_para.add_run(text)
                    run.font.name = font_name
                    run.font.size = Pt(14)
    
    # 处理表格（如果有）
    for table in doc.tables:
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        new_table.style = 'Light Grid Accent 1'
        
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_cell = new_table.rows[i].cells[j]
                new_cell.text = cell.text
                # 设置表格字体
                for paragraph in new_cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = Pt(12)
    
    # 保存新文档
    new_doc.save(output_path)
    print(f"✅ 处理完成！已保存到: {output_path}")
    print(f"📊 原段落数: {len(doc.paragraphs)}, 增强后段落数: {len(new_doc.paragraphs)}")
    return True

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("使用方法: python3 enrich_docx.py <输入文件路径> [输出文件路径]")
        print("示例: python3 enrich_docx.py doc/项目宣传2_公众号版_优化版_原创版.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    try:
        enrich_docx(input_file, output_file)
    except ImportError:
        print("错误：需要安装python-docx库")
        print("请运行: pip install python-docx")
        sys.exit(1)
    except Exception as e:
        print(f"处理文件时出错: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

